"""File parsing utilities for PDF, DOCX, CSV, XLSX, JSON, SQL, YAML, Notebook."""
import os
import json
import csv
from io import BytesIO, StringIO
from typing import Dict, List, Optional, Union
from pathlib import Path

from config.logging_config import get_logger

logger = get_logger("file_parser")

class FileParser:
    """Parse various file types into text content."""

    @staticmethod
    def parse_file(file_path: Union[str, bytes], file_name: str = None) -> Dict:
        """Parse a file based on its extension."""
        if isinstance(file_path, str):
            ext = Path(file_path).suffix.lower()
            file_name = file_name or Path(file_path).name
            with open(file_path, 'rb') as f:
                content = f.read()
        else:
            content = file_path
            ext = Path(file_name).suffix.lower() if file_name else '.txt'

        parsers = {
            '.pdf': FileParser.parse_pdf,
            '.docx': FileParser.parse_docx,
            '.doc': FileParser.parse_docx,
            '.csv': FileParser.parse_csv,
            '.xlsx': FileParser.parse_excel,
            '.xls': FileParser.parse_excel,
            '.json': FileParser.parse_json,
            '.sql': FileParser.parse_text,
            '.py': FileParser.parse_text,
            '.yaml': FileParser.parse_text,
            '.yml': FileParser.parse_text,
            '.ipynb': FileParser.parse_notebook,
            '.txt': FileParser.parse_text,
            '.md': FileParser.parse_text,
        }

        parser = parsers.get(ext, FileParser.parse_text)

        try:
            result = parser(content, file_name)
            logger.info(f"Parsed {file_name} ({ext})")
            return result
        except Exception as e:
            logger.error(f"Failed to parse {file_name}: {e}")
            return {"text": "", "metadata": {"error": str(e), "file_name": file_name}}

    @staticmethod
    def parse_pdf(content: bytes, file_name: str = None) -> Dict:
        """Parse PDF file."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "
"
            return {
                "text": text,
                "metadata": {
                    "file_name": file_name,
                    "pages": len(reader.pages),
                    "type": "pdf"
                }
            }
        except ImportError:
            logger.warning("pypdf not installed, returning raw content")
            return {"text": content.decode('utf-8', errors='ignore'), "metadata": {"file_name": file_name, "type": "pdf"}}

    @staticmethod
    def parse_docx(content: bytes, file_name: str = None) -> Dict:
        """Parse DOCX file."""
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            text = "
".join([para.text for para in doc.paragraphs])
            return {
                "text": text,
                "metadata": {
                    "file_name": file_name,
                    "paragraphs": len(doc.paragraphs),
                    "type": "docx"
                }
            }
        except ImportError:
            logger.warning("python-docx not installed, returning raw content")
            return {"text": content.decode('utf-8', errors='ignore'), "metadata": {"file_name": file_name, "type": "docx"}}

    @staticmethod
    def parse_csv(content: bytes, file_name: str = None) -> Dict:
        """Parse CSV file."""
        text = content.decode('utf-8', errors='ignore')
        reader = csv.DictReader(StringIO(text))
        rows = list(reader)

        # Create a structured text representation
        structured_text = f"CSV File: {file_name}
"
        structured_text += f"Columns: {', '.join(reader.fieldnames)}

"
        structured_text += f"Total Rows: {len(rows)}

"
        structured_text += "Sample Data (first 10 rows):
"
        for i, row in enumerate(rows[:10]):
            structured_text += f"Row {i+1}: {json.dumps(row)}
"

        return {
            "text": structured_text,
            "metadata": {
                "file_name": file_name,
                "columns": reader.fieldnames,
                "row_count": len(rows),
                "type": "csv"
            }
        }

    @staticmethod
    def parse_excel(content: bytes, file_name: str = None) -> Dict:
        """Parse Excel file."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(content))
            text = f"Excel File: {file_name}
"
            text += f"Sheets: {wb.sheetnames}

"

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"--- Sheet: {sheet_name} ---
"
                for row in sheet.iter_rows(values_only=True):
                    text += ", ".join([str(cell) if cell is not None else "" for cell in row]) + "
"
                text += "
"

            return {
                "text": text,
                "metadata": {
                    "file_name": file_name,
                    "sheets": wb.sheetnames,
                    "type": "excel"
                }
            }
        except ImportError:
            logger.warning("openpyxl not installed, returning raw content")
            return {"text": content.decode('utf-8', errors='ignore'), "metadata": {"file_name": file_name, "type": "excel"}}

    @staticmethod
    def parse_json(content: bytes, file_name: str = None) -> Dict:
        """Parse JSON file."""
        text = content.decode('utf-8', errors='ignore')
        try:
            data = json.loads(text)
            structured_text = f"JSON File: {file_name}
"
            structured_text += f"Type: {type(data).__name__}

"
            structured_text += json.dumps(data, indent=2)[:5000]  # Limit size
            return {
                "text": structured_text,
                "metadata": {
                    "file_name": file_name,
                    "type": "json"
                }
            }
        except json.JSONDecodeError:
            return {"text": text, "metadata": {"file_name": file_name, "type": "json", "parse_error": True}}

    @staticmethod
    def parse_text(content: bytes, file_name: str = None) -> Dict:
        """Parse plain text files."""
        text = content.decode('utf-8', errors='ignore')
        return {
            "text": text,
            "metadata": {
                "file_name": file_name,
                "type": "text",
                "length": len(text)
            }
        }

    @staticmethod
    def parse_notebook(content: bytes, file_name: str = None) -> Dict:
        """Parse Jupyter notebook."""
        text = content.decode('utf-8', errors='ignore')
        try:
            nb = json.loads(text)
            structured_text = f"Notebook: {file_name}
"
            structured_text += f"Cells: {len(nb.get('cells', []))}

"

            for i, cell in enumerate(nb.get('cells', [])):
                cell_type = cell.get('cell_type', 'unknown')
                source = ''.join(cell.get('source', []))
                structured_text += f"--- Cell {i+1} ({cell_type}) ---
"
                structured_text += source[:1000] + "

"  # Limit per cell

            return {
                "text": structured_text,
                "metadata": {
                    "file_name": file_name,
                    "cell_count": len(nb.get('cells', [])),
                    "type": "notebook"
                }
            }
        except json.JSONDecodeError:
            return {"text": text, "metadata": {"file_name": file_name, "type": "notebook", "parse_error": True}}


def parse_uploaded_file(uploaded_file) -> Dict:
    """Parse an uploaded file from Streamlit."""
    return FileParser.parse_file(uploaded_file.getvalue(), uploaded_file.name)
